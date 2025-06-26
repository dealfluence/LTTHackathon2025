import pdfplumber
import docx
from pathlib import Path
from typing import Dict, Any, List
from ..interfaces.document_source import DocumentSource


class LocalFileSource(DocumentSource):

    def __init__(self):
        self.supported_formats = ["pdf", "docx", "txt"]

    def get_supported_formats(self) -> List[str]:
        return self.supported_formats

    def validate_source(self, source_path: str) -> bool:
        try:
            path = Path(source_path)
            return (
                path.exists()
                and path.is_file()
                and path.suffix.lower().lstrip(".") in self.supported_formats
            )
        except Exception:
            return False

    async def load_document(self, source_path: str) -> Dict[str, Any]:
        if not self.validate_source(source_path):
            raise ValueError(f"Invalid source path: {source_path}")

        path = Path(source_path)
        file_extension = path.suffix.lower()

        if file_extension == ".pdf":
            return await self._process_pdf(path)
        elif file_extension == ".docx":
            return await self._process_docx(path)
        elif file_extension == ".txt":
            return await self._process_txt(path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    async def _process_pdf(self, path: Path) -> Dict[str, Any]:
        try:
            text_content = ""
            metadata = {"pages": 0, "file_type": "pdf"}

            with pdfplumber.open(path) as pdf:
                metadata["pages"] = len(pdf.pages)
                for page in pdf.pages:
                    if page.extract_text():
                        text_content += page.extract_text() + "\n"

            return {
                "content": text_content.strip(),
                "metadata": metadata,
                "source_path": str(path),
                "filename": path.name,
            }
        except Exception as e:
            raise Exception(f"Error processing PDF: {e}")

    async def _process_docx(self, path: Path) -> Dict[str, Any]:
        try:
            doc = docx.Document(path)
            text_content = ""

            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"

            metadata = {"paragraphs": len(doc.paragraphs), "file_type": "docx"}

            return {
                "content": text_content.strip(),
                "metadata": metadata,
                "source_path": str(path),
                "filename": path.name,
            }
        except Exception as e:
            raise Exception(f"Error processing DOCX: {e}")

    async def _process_txt(self, path: Path) -> Dict[str, Any]:
        try:
            with open(path, "r", encoding="utf-8") as file:
                content = file.read()

            metadata = {"characters": len(content), "file_type": "txt"}

            return {
                "content": content,
                "metadata": metadata,
                "source_path": str(path),
                "filename": path.name,
            }
        except Exception as e:
            raise Exception(f"Error processing TXT: {e}")
